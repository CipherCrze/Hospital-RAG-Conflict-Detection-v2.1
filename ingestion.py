"""
Document Ingestion Pipeline
- Load PDFs, DOCX, and TXT files
- Chunk documents with metadata
- Store embeddings in ChromaDB
"""
import os
import sys
import io
import re

# Fix Windows console encoding for tqdm/transformers Unicode output
if sys.stdout and hasattr(sys.stdout, 'buffer'):
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass
if sys.stderr and hasattr(sys.stderr, 'buffer'):
    try:
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.embeddings import Embeddings
from sentence_transformers import SentenceTransformer as _ST
from config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    EMBEDDING_MODEL,
    CHROMA_PERSIST_DIR,
    COLLECTION_NAME,
    SUPPORTED_EXTENSIONS,
)


class _STEmbeddings(Embeddings):
    """LangChain-compatible wrapper around SentenceTransformer (no async client)."""
    def __init__(self, model_name: str):
        self._m = _ST(model_name)
    def embed_documents(self, texts):
        return self._m.encode(texts, show_progress_bar=False).tolist()
    def embed_query(self, text):
        return self._m.encode([text], show_progress_bar=False)[0].tolist()


# Module-level singleton — SentenceTransformer loads once per process.
_EMBEDDINGS_SINGLETON: _STEmbeddings | None = None


def _embeddings() -> _STEmbeddings:
    global _EMBEDDINGS_SINGLETON
    if _EMBEDDINGS_SINGLETON is None:
        _EMBEDDINGS_SINGLETON = _STEmbeddings(EMBEDDING_MODEL)
    return _EMBEDDINGS_SINGLETON


def load_txt(filepath: str) -> list[dict]:
    """Load a plain text file and return pages."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [{"text": text, "page": 1}]


def load_pdf(filepath: str) -> list[dict]:
    """Load a PDF file and return pages with text."""
    pages = []
    doc = fitz.open(filepath)
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append({"text": text, "page": i + 1})
    doc.close()
    return pages


def load_docx(filepath: str) -> list[dict]:
    """Load a DOCX file and return as a single page."""
    doc = DocxDocument(filepath)
    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [{"text": text, "page": 1}]


def load_document(filepath: str) -> list[dict]:
    """Load a document based on its extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        return load_txt(filepath)
    elif ext == ".pdf":
        return load_pdf(filepath)
    elif ext == ".docx":
        return load_docx(filepath)
    else:
        print(f"  [WARN] Unsupported file type: {ext}")
        return []


def extract_department(filename: str, text: str) -> str:
    """Infer department from filename and content."""
    fname_lower = filename.lower()
    text_lower = text[:500].lower()

    if "emergency" in fname_lower or "emergency" in text_lower:
        return "Emergency"
    elif "surg" in fname_lower or "surgical" in text_lower:
        return "Surgery"
    elif "nurs" in fname_lower or "nursing" in text_lower:
        return "Nursing"
    elif "outpatient" in fname_lower or "outpatient" in text_lower:
        return "Outpatient"
    elif "infection" in fname_lower or "infection" in text_lower:
        return "Infection Control"
    elif "financial" in fname_lower or "budget" in fname_lower or "finance" in text_lower:
        return "Finance"
    elif "board" in fname_lower or "board" in text_lower:
        return "Administration"
    elif "satisfaction" in fname_lower or "survey" in fname_lower:
        return "Quality"
    elif "complaint" in fname_lower or "grievance" in fname_lower:
        return "Patient Relations"
    elif "turnover" in fname_lower or "staff" in fname_lower or "human resources" in text_lower:
        return "Human Resources"
    else:
        return "General"


def chunk_documents(docs_dir: str) -> list:
    """
    Load all documents from a directory, chunk them, and return
    LangChain Document objects with metadata.
    """
    from langchain_core.documents import Document

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    all_chunks = []
    files = sorted(os.listdir(docs_dir))

    for filename in files:
        filepath = os.path.join(docs_dir, filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        print(f"  [LOAD] Loading: {filename}")
        pages = load_document(filepath)

        if not pages:
            print(f"  [WARN] No content extracted from {filename}")
            continue

        for page_data in pages:
            text = page_data["text"]
            page_num = page_data["page"]
            department = extract_department(filename, text)

            chunks = splitter.split_text(text)

            for i, chunk_text in enumerate(chunks):
                doc = Document(
                    page_content=chunk_text,
                    metadata={
                        "source": filename,
                        "doc_id": os.path.splitext(filename)[0],
                        "page": page_num,
                        "chunk_id": i,
                        "department": department,
                        "quarter": "Q1 2025",
                        "doc_type": ext.replace(".", "").upper(),
                    },
                )
                all_chunks.append(doc)

    print(f"\n  [OK] Total chunks created: {len(all_chunks)}")
    return all_chunks


def create_vector_store(chunks: list, persist_dir: str = CHROMA_PERSIST_DIR):
    """
    Create a ChromaDB vector store from document chunks.

    Returns:
        Chroma vector store instance.
    """
    print(f"\n[EMBED] Creating embeddings with {EMBEDDING_MODEL}...")
    embeddings = _embeddings()

    # Remove existing DB if present
    if os.path.exists(persist_dir):
        import shutil
        shutil.rmtree(persist_dir)
        print("  [CLEAR] Cleared existing vector store")

    print(f"  [STORE] Storing {len(chunks)} chunks in ChromaDB...")
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=COLLECTION_NAME,
        persist_directory=persist_dir,
    )

    print(f"  [OK] Vector store created at: {persist_dir}")
    return vector_store


def load_vector_store(persist_dir: str = CHROMA_PERSIST_DIR):
    """Load an existing ChromaDB vector store (embedding model cached at module level)."""
    return Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=persist_dir,
        embedding_function=_embeddings(),
    )


def ingest_documents(docs_dir: str = None):
    """
    Full ingestion pipeline: load → chunk → embed → store.
    """
    from config import SAMPLE_DOCS_DIR

    if docs_dir is None:
        docs_dir = SAMPLE_DOCS_DIR

    print(f"[INGEST] Ingesting documents from: {docs_dir}")
    print("=" * 50)

    chunks = chunk_documents(docs_dir)

    if not chunks:
        print("[ERROR] No chunks created. Check your documents directory.")
        return None

    vector_store = create_vector_store(chunks)
    return vector_store


if __name__ == "__main__":
    store = ingest_documents()
    if store:
        # Quick verification
        collection = store._collection
        print(f"\n[VERIFY] Verification:")
        print(f"  Collection: {COLLECTION_NAME}")
        print(f"  Total documents: {collection.count()}")
        print("  [OK] Ingestion complete!")
